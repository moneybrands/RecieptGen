from docx import Document
from docx2pdf import convert
from datetime import datetime
from random import randint
from num2words import num2words
import os
from openpyxl import load_workbook 


def edit(AC,Amount,Name,Date,Arrears="0.00"):
    doc = Document("receipt.docx")
    for col in [doc.tables[0].columns[1],doc.tables[0].columns[3]]:
        for c in col.cells:
            for paragraph in c.paragraphs:
                for cell in paragraph.runs:
                    
                    cell.text = cell.text.replace("333", str(randint(100, 999)))
                    cell.text = cell.text.replace("ACNumber", str(AC))
                    cell.text = cell.text.replace("Amount", " "+str(Amount))
                    cell.text = cell.text.replace("Words", f'{num2words(Amount).replace(" and "," ").title()} Only')
                    cell.text = cell.text.replace("sandarbh", str(randint(10000, 99900)))
                    cell.text = cell.text.replace("date", Date)
                    cell.text = cell.text.replace("xyz", Name.upper())
                    cell.text = cell.text.replace("bakaya", str(float(Arrears)))  #if given

    doc.save(rf'Temp/{AC}.docx')
    convert(rf'Temp/{AC}.docx',f"Temp/{Name}_{AC}.pdf")


def getDetail():

    AC=input("\033[1m" + "Account Number Kya Hai:- " + "\033[0m")
    if len(AC)!=10:
        print("\033[1;31m" + "\n          Account Number Ki Length 10 Honi Chahiye\n" + "\033[0m")
        getDetail()


    Name=str(input("\033[1m"+"Name Batao:- "+"\033[0m"))
    if any(chr.isdigit() for chr in Name):
        print("\033[1;31m" + "\n          Name Me Number Kaise Aa Gya? Try Again.../n" + "\033[0m")
        getDetail()

    
    Amount=str(input("\033[1m"+"Amount Kitna Hai:- "+"\033[0m"))
        
    Date=datetime.now().strftime("%d-%m-%y %H:%M:%S")

    if "," in Amount:
        amount=Amount.split(",")
        edit(AC,float(amount[0]),Name,Date,Arrears=float(amount[1]))

    else:
        edit(AC,Amount,Name,Date)

    return AC, Name, Amount, Date



    
def final(AC, Name, Amount, Date):
    #Updating Excel
    
    row_data = [AC, Name,Date,Amount.split(",")[0]]
    workbook = load_workbook('data.xlsx')
    sheet = workbook.active.append(row_data)
    workbook.save('data.xlsx')
    os.startfile(f"Temp\{Name}_{AC}.pdf")
    
    print("\033[1;32m" + "\nReciept Generated! Enjoy :)\n" + "\033[0m")
    
    #Deleting Files
    folder=os.path.join(os.getcwd(), "Temp")
    for file in os.listdir(folder):
        if file.endswith(".docx"):
            os.remove(os.path.join(folder, file))
        
    
    print("\033[1;34m" + "\n\n-------------------------------------------------------------------\n\n" + "\033[0m")


while True:
    try:
        AC, Name, Amount, Date = getDetail()
        final(AC, Name, Amount, Date)
    except Exception as e:
        print("Error")
        print(e)
        #print("\033[1;31m",e,"\033[0m")
        getDetail()
